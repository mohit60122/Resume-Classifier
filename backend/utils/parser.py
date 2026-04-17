"""
Resume Parser - Extract text from PDF and DOCX files.
"""
import os
from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text content from a PDF file."""
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    """Extract text content from a DOCX file."""
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text.strip()


def extract_text(file_path: str) -> str:
    """
    Auto-detect file type and extract text.
    Supports: .pdf, .docx
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported: .pdf, .docx")
